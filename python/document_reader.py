"""
document_reader.py
==================
Smart Document Management System — Week 6 | Student A
------------------------------------------------------
Responsibilities:
    - Extract raw text from PDF, DOCX, and TXT documents
    - Pull document metadata (author, title, page count, timestamps, size)
    - Feed extracted text into TextPreprocessor to produce clean tokens
    - Return one fully-structured dict consumed by api.py / Laravel

Supported formats : .pdf   .docx   .txt
Python            : 3.8+
Required packages : pdfplumber  pypdf  python-docx
                    (already installed in previous step)

NO external binaries required. No Tesseract. No extra apps.
"""

from __future__ import annotations

import hashlib
import logging
import os
import re
import sys
import unicodedata
from datetime import datetime
from pathlib import Path
from typing import Any

# ── PDF ───────────────────────────────────────────────────────────────────────
try:
    import pdfplumber
    _PDF_PLUMBER = True
except ImportError:
    _PDF_PLUMBER = False

try:
    from pypdf import PdfReader
    _PYPDF = True
except ImportError:
    _PYPDF = False

# ── DOCX ──────────────────────────────────────────────────────────────────────
try:
    import docx as _docx
    _DOCX = True
except ImportError:
    _DOCX = False

# ── TextPreprocessor (your existing file) ─────────────────────────────────────
try:
    from text_preprocessing import TextPreprocessor
    _PREPROCESSOR = True
except ImportError:
    _PREPROCESSOR = False

# ── Logging ───────────────────────────────────────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s  [%(levelname)s]  %(name)s — %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
)
logger = logging.getLogger("document_reader")

# ── Constants ─────────────────────────────────────────────────────────────────
SUPPORTED_EXTENSIONS   = {".pdf", ".docx", ".txt"}
MAX_TEXT_PREVIEW_CHARS = 500


# =============================================================================
#  HELPERS
# =============================================================================

def _normalise_text(raw: str) -> str:
    """Unicode-normalise, strip control characters, collapse excess whitespace."""
    text = unicodedata.normalize("NFC", raw)
    text = re.sub(r"[^\S\n\t ]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = "\n".join(line.strip() for line in text.splitlines())
    return text.strip()


def _file_hash(path: str) -> str:
    """SHA-256 fingerprint — used by Laravel for deduplication before DB insert."""
    sha = hashlib.sha256()
    with open(path, "rb") as fh:
        for chunk in iter(lambda: fh.read(65536), b""):
            sha.update(chunk)
    return sha.hexdigest()


def _file_metadata(path: str) -> dict[str, Any]:
    """File-system level metadata common to every format."""
    stat = os.stat(path)
    return {
        "file_name":       os.path.basename(path),
        "file_size_kb":    round(stat.st_size / 1024, 2),
        "file_size_bytes": stat.st_size,
        "last_modified":   datetime.fromtimestamp(stat.st_mtime).isoformat(),
        "sha256":          _file_hash(path),
    }


# =============================================================================
#  PDF
# =============================================================================

def _pdf_metadata(path: str) -> dict[str, Any]:
    """Read embedded PDF metadata (author, title, dates, page count)."""
    meta: dict[str, Any] = {}
    if not _PYPDF:
        return meta
    try:
        reader = PdfReader(path)
        info   = reader.metadata or {}
        meta["title"]      = str(info.get("/Title",    "")).strip() or None
        meta["author"]     = str(info.get("/Author",   "")).strip() or None
        meta["subject"]    = str(info.get("/Subject",  "")).strip() or None
        meta["creator"]    = str(info.get("/Creator",  "")).strip() or None
        meta["producer"]   = str(info.get("/Producer", "")).strip() or None
        meta["page_count"] = len(reader.pages)
        for key, pdf_key in [("created_at", "/CreationDate"),
                              ("modified_at", "/ModDate")]:
            raw = str(info.get(pdf_key, "")).strip()
            if raw.startswith("D:"):
                try:
                    meta[key] = datetime.strptime(raw[2:16], "%Y%m%d%H%M%S").isoformat()
                except ValueError:
                    meta[key] = raw
            else:
                meta[key] = raw or None
    except Exception as exc:
        logger.warning("PDF metadata error: %s", exc)
    return meta


def _extract_pdf(path: str) -> tuple[str, dict[str, Any]]:
    """
    Extract text page-by-page with pdfplumber.
    Each page is labelled so downstream search can reference exact pages.
    """
    if not _PDF_PLUMBER:
        raise EnvironmentError(
            "pdfplumber not installed. Run: pip install pdfplumber"
        )
    pages: list[str] = []
    with pdfplumber.open(path) as pdf:
        total = len(pdf.pages)
        for num, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            pages.append(f"[Page {num}/{total}]\n{text}")
            logger.info("PDF page %d/%d extracted.", num, total)

    full_text = "\n\n".join(pages)
    meta      = _pdf_metadata(path)
    return _normalise_text(full_text), meta


# =============================================================================
#  DOCX
# =============================================================================

def _docx_metadata(path: str) -> dict[str, Any]:
    """Read DOCX core properties."""
    meta: dict[str, Any] = {}
    if not _DOCX:
        return meta
    try:
        doc   = _docx.Document(path)
        props = doc.core_properties
        meta["title"]       = props.title  or None
        meta["author"]      = props.author or None
        meta["subject"]     = props.subject or None
        meta["created_at"]  = props.created.isoformat()  if props.created  else None
        meta["modified_at"] = props.modified.isoformat() if props.modified else None
        meta["word_count"]  = sum(
            len(p.text.split()) for p in doc.paragraphs if p.text.strip()
        )
    except Exception as exc:
        logger.warning("DOCX metadata error: %s", exc)
    return meta


def _extract_docx(path: str) -> tuple[str, dict[str, Any]]:
    """
    Extract text from DOCX.
    Preserves heading hierarchy and includes table cell content.
    """
    if not _DOCX:
        raise EnvironmentError(
            "python-docx not installed. Run: pip install python-docx"
        )
    doc      = _docx.Document(path)
    sections: list[str] = []

    for para in doc.paragraphs:
        stripped = para.text.strip()
        if not stripped:
            continue
        if para.style.name.startswith("Heading"):
            level  = re.search(r"\d", para.style.name)
            prefix = "#" * int(level.group()) if level else "##"
            sections.append(f"{prefix} {stripped}")
        else:
            sections.append(stripped)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                c.text.strip() for c in row.cells if c.text.strip()
            )
            if row_text:
                sections.append(row_text)

    meta = _docx_metadata(path)
    return _normalise_text("\n".join(sections)), meta


# =============================================================================
#  TXT
# =============================================================================

def _extract_txt(path: str) -> tuple[str, dict[str, Any]]:
    """Read plain text; auto-detects encoding."""
    for enc in ("utf-8", "utf-8-sig", "windows-1252", "latin-1"):
        try:
            with open(path, "r", encoding=enc) as fh:
                text = fh.read()
            meta: dict[str, Any] = {
                "encoding":   enc,
                "line_count": text.count("\n"),
                "word_count": len(text.split()),
            }
            logger.info("TXT decoded as %s.", enc)
            return _normalise_text(text), meta
        except (UnicodeDecodeError, LookupError):
            continue
    raise ValueError(f"Could not decode {path} — unknown encoding.")


# =============================================================================
#  PUBLIC ENTRY POINT
# =============================================================================

def read_document(file_path: str) -> dict[str, Any]:
    """
    Extract text and metadata from a document file.
    This is the function called by api.py and by Laravel via the Python API.

    Parameters
    ----------
    file_path : str
        Path to the document (.pdf / .docx / .txt)

    Returns
    -------
    dict
        success           bool
        file_name         str
        file_type         str
        raw_text          str       full extracted text
        raw_length        int       character count
        preview           str       first 500 characters
        metadata          dict      author, title, page_count, dates, sha256 …
        tokens            list      preprocessed tokens from TextPreprocessor
        token_count       int
        cleaned_text      str       preprocessed string
        extraction_method str
        processed_at      str       ISO-8601 UTC timestamp
        error             str|None
    """
    logger.info("Reading document: %s", file_path)

    # Validate
    if not os.path.exists(file_path):
        return {"success": False, "error": f"File not found: {file_path}"}

    ext = Path(file_path).suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        return {
            "success": False,
            "error": (
                f"Unsupported format '{ext}'. "
                f"Accepted: {sorted(SUPPORTED_EXTENSIONS)}"
            ),
        }

    fs_meta = _file_metadata(file_path)

    # Extract
    try:
        if ext == ".pdf":
            raw_text, fmt_meta = _extract_pdf(file_path)
            method = "pdfplumber"
        elif ext == ".docx":
            raw_text, fmt_meta = _extract_docx(file_path)
            method = "python-docx"
        else:
            raw_text, fmt_meta = _extract_txt(file_path)
            method = "plaintext"
    except Exception as exc:
        logger.exception("Extraction failed.")
        return {"success": False, "error": str(exc),
                "file_name": fs_meta["file_name"]}

    if not raw_text.strip():
        return {
            "success": False,
            "error": "No readable text found in document.",
            "file_name": fs_meta["file_name"],
        }

    # Preprocess via your TextPreprocessor
    tokens: list[str] = []
    cleaned_text      = ""
    if _PREPROCESSOR:
        try:
            pp           = TextPreprocessor()
            tokens       = pp.preprocess(raw_text)
            cleaned_text = pp.preprocess(raw_text, return_string=True)
        except Exception as exc:
            logger.warning("TextPreprocessor error (non-fatal): %s", exc)
    else:
        logger.warning("text_preprocessing.py not found — skipping preprocessing.")

    return {
        "success":            True,
        "file_name":          fs_meta["file_name"],
        "file_type":          ext,
        "raw_text":           raw_text,
        "raw_length":         len(raw_text),
        "preview":            raw_text[:MAX_TEXT_PREVIEW_CHARS],
        "metadata":           {**fs_meta, **fmt_meta},
        "tokens":             tokens,
        "token_count":        len(tokens),
        "cleaned_text":       cleaned_text,
        "extraction_method":  method,
        "processed_at":       datetime.utcnow().isoformat() + "Z",
        "error":              None,
    }


# =============================================================================
#  CLI
#  python document_reader.py            → self-test
#  python document_reader.py file.pdf   → process a real file
# =============================================================================

def _self_test() -> None:
    import tempfile
    sample = (
        "Smart Document Management System\n\n"
        "This system integrates Artificial Intelligence with a RESTful API "
        "to deliver intelligent document search and retrieval.\n\n"
        "Key features:\n"
        "  - AI-powered semantic search\n"
        "  - Contextual document ranking\n"
        "  - Support for PDF, DOCX, and TXT formats\n"
        "  - Scalable design via Laravel + Python microservice\n"
    )
    with tempfile.NamedTemporaryFile(
        mode="w", suffix=".txt", delete=False, encoding="utf-8"
    ) as tmp:
        tmp.write(sample)
        tmp_path = tmp.name
    try:
        _print_result(read_document(tmp_path))
    finally:
        os.unlink(tmp_path)


def _print_result(r: dict[str, Any]) -> None:
    SEP  = "─" * 66
    SEP2 = "═" * 66
    if not r["success"]:
        print(f"\n{SEP}\n  ERROR: {r.get('error')}\n{SEP}\n")
        return
    print(f"\n{SEP2}")
    print("  DOCUMENT READER  —  Extraction Report")
    print(f"{SEP2}")
    print(f"  {'File':<22}: {r['file_name']}")
    print(f"  {'Type':<22}: {r['file_type']}")
    print(f"  {'Method':<22}: {r['extraction_method']}")
    print(f"  {'Processed (UTC)':<22}: {r['processed_at']}")
    print(SEP)
    print("  METADATA")
    print(SEP)
    for k, v in r["metadata"].items():
        if v not in (None, "", [], {}):
            print(f"  {k:<22}: {v}")
    print(SEP)
    print("  TEXT STATISTICS")
    print(SEP)
    print(f"  {'Raw characters':<22}: {r['raw_length']:,}")
    print(f"  {'Token count':<22}: {r['token_count']:,}")
    print(SEP)
    print("  PREVIEW  (first 500 chars)")
    print(SEP)
    print(r["preview"])
    if r["cleaned_text"]:
        print(SEP)
        print("  CLEANED TEXT")
        print(SEP)
        print(r["cleaned_text"][:500])
        print(SEP)
        print("  TOKENS  (first 30)")
        print(SEP)
        print(r["tokens"][:30])
    print(SEP2 + "\n")


if __name__ == "__main__":
    if len(sys.argv) > 1:
        _print_result(read_document(sys.argv[1]))
    else:
        logger.info("No file given — running self-test.")
        _self_test()
